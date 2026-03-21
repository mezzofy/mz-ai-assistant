"""
DOCX Tool — Generate and read Word documents.

Tools provided:
    create_docx  — Generate a Mezzofy-branded Word document
    read_docx    — Extract text from an uploaded DOCX file

Uses python-docx for document generation. Mezzofy brand colors:
    Orange: #f97316  |  Black: #000000

Output files saved to the configured artifact storage directory.
"""

import logging
import uuid
from pathlib import Path
from typing import Optional

from app.context.artifact_manager import get_user_artifacts_dir, get_dept_artifacts_dir, get_company_artifacts_dir
from app.core.user_context import get_user_dept, get_user_email
from app.tools.base_tool import BaseTool

logger = logging.getLogger("mezzofy.tools.docx")

# Mezzofy brand orange
_ORANGE_HEX = "F97316"


def _get_logo_path(filename: str):
    """
    Resolve a logo file from knowledge/brand/.
    Returns Path if it exists, None otherwise.
    Path: 4 levels up from app/tools/document/ → server root → knowledge/brand/
    """
    logo_path = Path(__file__).parent.parent.parent.parent / "knowledge" / "brand" / filename
    return logo_path if logo_path.exists() else None


def _get_artifact_dir(config: dict) -> Path:
    base = config.get("storage", {}).get("local_path", "/data/artifacts")
    path = Path(base) / "docx"
    path.mkdir(parents=True, exist_ok=True)
    return path


class DocxOps(BaseTool):
    """Word document generation and reading."""

    def __init__(self, config: dict):
        super().__init__(config)
        self._artifact_dir = _get_artifact_dir(config)

    def _resolve_output_dir(self, storage_scope: str = "user") -> Path:
        """Return output dir based on storage scope and user context."""
        dept = get_user_dept()
        email = get_user_email()
        if storage_scope == "company":
            return get_company_artifacts_dir()
        if storage_scope == "department" and dept:
            return get_dept_artifacts_dir(dept)
        if email:
            return get_user_artifacts_dir(dept, email)
        return self._artifact_dir

    def get_tools(self) -> list[dict]:
        return [
            {
                "name": "create_docx",
                "description": (
                    "Generate a Mezzofy-branded Word document (.docx) from structured content. "
                    "Supports headings, paragraphs, tables, and bullet lists. "
                    "Returns the file path of the generated document."
                ),
                "parameters": {
                    "type": "object",
                    "properties": {
                        "title": {
                            "type": "string",
                            "description": "Document title shown at the top.",
                        },
                        "sections": {
                            "type": "array",
                            "description": "Document sections. Each section can be a heading, paragraph, table, or list.",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "type": {
                                        "type": "string",
                                        "description": "Content type.",
                                        "enum": ["heading1", "heading2", "heading3", "paragraph", "table", "list"],
                                    },
                                    "content": {
                                        "type": "string",
                                        "description": (
                                            "Text content for heading, paragraph, or list. "
                                            "For list type: one item per line."
                                        ),
                                    },
                                    "table_data": {
                                        "type": "array",
                                        "description": (
                                            "Table rows for 'table' type. First row is header. "
                                            "Each row is a list of cell strings."
                                        ),
                                        "items": {
                                            "type": "array",
                                            "items": {"type": "string"},
                                        },
                                    },
                                },
                                "required": ["type"],
                            },
                        },
                        "author": {
                            "type": "string",
                            "description": "Author name shown in document metadata.",
                        },
                        "filename": {
                            "type": "string",
                            "description": "Output filename (without extension). Auto-generated if omitted.",
                        },
                        "storage_scope": {
                            "type": "string",
                            "description": (
                                "Where to save the file. 'user' = personal folder (default), "
                                "'department' = shared department folder, "
                                "'company' = company-wide public folder (management only)."
                            ),
                            "enum": ["user", "department", "company"],
                            "default": "user",
                        },
                    },
                    "required": ["title", "sections"],
                },
                "handler": self._create_docx,
            },
            {
                "name": "read_docx",
                "description": (
                    "Extract text content from an uploaded Word document (.docx). "
                    "Returns structured content including headings, paragraphs, and tables."
                ),
                "parameters": {
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": "Absolute path to the DOCX file to read.",
                        },
                    },
                    "required": ["file_path"],
                },
                "handler": self._read_docx,
            },
        ]

    async def _create_docx(
        self,
        title: str,
        sections: list[dict],
        author: Optional[str] = None,
        filename: Optional[str] = None,
        storage_scope: str = "user",
    ) -> dict:
        """Generate a branded Word document."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            return self._err("python-docx is not installed. Run: pip install python-docx")

        if not filename:
            safe = "".join(c if c.isalnum() or c in "-_ " else "" for c in title)
            safe = safe.replace(" ", "_")[:40]
            filename = f"{safe}_{uuid.uuid4().hex[:8]}"

        output_path = self._resolve_output_dir(storage_scope) / f"{filename}.docx"

        doc = Document()

        # Document properties
        if author:
            doc.core_properties.author = author
        doc.core_properties.company = "Mezzofy"

        # Customize built-in heading styles with Mezzofy orange
        orange = RGBColor(0xF9, 0x73, 0x16)
        for style_name in ("Heading 1", "Heading 2", "Heading 3"):
            style = doc.styles[style_name]
            style.font.color.rgb = orange

        # --- Header ---
        logo_path = _get_logo_path("logo_dark.png")  # white background → dark logo
        section = doc.sections[0]
        header = section.header
        hdr_para = header.paragraphs[0]
        hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hdr_run = hdr_para.add_run()
        if logo_path:
            hdr_run.add_picture(str(logo_path), width=Inches(1.2))
        else:
            hdr_run.text = "Mezzofy AI Assistant"
            hdr_run.font.color.rgb = RGBColor(0xF9, 0x73, 0x16)
            hdr_run.font.bold = True

        # --- Title ---
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in title_para.runs:
            run.font.color.rgb = orange

        doc.add_paragraph()  # spacing

        # --- Sections ---
        for item in sections:
            item_type = item.get("type", "paragraph")
            content = item.get("content", "")

            if item_type == "heading1":
                doc.add_heading(content, level=1)

            elif item_type == "heading2":
                doc.add_heading(content, level=2)

            elif item_type == "heading3":
                doc.add_heading(content, level=3)

            elif item_type == "paragraph":
                doc.add_paragraph(content)

            elif item_type == "list":
                lines = content.strip().split("\n") if content else []
                for line in lines:
                    line = line.lstrip("•-* ").strip()
                    if line:
                        doc.add_paragraph(line, style="List Bullet")

            elif item_type == "table":
                table_data = item.get("table_data", [])
                if not table_data:
                    continue

                rows = len(table_data)
                cols = max(len(r) for r in table_data)
                table = doc.add_table(rows=rows, cols=cols)
                table.style = "Table Grid"

                for r_idx, row in enumerate(table_data):
                    for c_idx, cell_text in enumerate(row):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = str(cell_text)
                        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                        if run and r_idx == 0:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            # Set header cell background orange
                            shading = OxmlElement("w:shd")
                            shading.set(qn("w:fill"), _ORANGE_HEX)
                            shading.set(qn("w:color"), "auto")
                            shading.set(qn("w:val"), "clear")
                            cell._tc.get_or_add_tcPr().append(shading)

                doc.add_paragraph()  # spacing after table

        # --- Footer ---
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Generated by Mezzofy AI Assistant"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer_para.runs:
            footer_para.runs[0].font.color.rgb = orange
            footer_para.runs[0].font.size = Pt(9)

        doc.save(str(output_path))
        file_size = output_path.stat().st_size

        logger.info(f"Created DOCX: {output_path}")
        return self._ok({
            "file_path": str(output_path),
            "filename": f"{filename}.docx",
            "size_bytes": file_size,
            "title": title,
            "sections_count": len(sections),
            "storage_scope": storage_scope,
            "department": get_user_dept(),
        })

    async def _read_docx(self, file_path: str) -> dict:
        """Extract text from a Word document."""
        import os
        if not os.path.exists(file_path):
            return self._err(f"File not found: {file_path}")

        try:
            from docx import Document
        except ImportError:
            return self._err("python-docx is not installed. Run: pip install python-docx")

        try:
            doc = Document(file_path)
            content = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else "Normal"
                content.append({
                    "type": style_name,
                    "text": text,
                })

            tables = []
            for t_idx, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                tables.append({
                    "table_index": t_idx,
                    "rows": rows,
                })

            return self._ok({
                "file_path": file_path,
                "paragraphs": content,
                "tables": tables,
                "paragraph_count": len(content),
                "table_count": len(tables),
            })

        except Exception as e:
            logger.error(f"Failed to read DOCX {file_path}: {e}")
            return self._err(str(e))
